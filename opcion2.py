import os
import pandas as pd
from docx import Document
import re

# Rutas a los archivos
plantilla = "/Users/paulalaorga/Development/cartas/Raw Data/Plantilla.docx"
datos_path = "/Users/paulalaorga/Development/cartas/Parsed/extract_data_output.csv"
output_folder = "/Users/paulalaorga/Development/cartas/Final"
os.makedirs(output_folder, exist_ok=True)

# Leer y preparar datos del archivo CSV
try:
    datos = pd.read_csv(datos_path, on_bad_lines="skip", delimiter="|")
    print(f"Archivo CSV cargado correctamente. Columnas encontradas: {list(datos.columns)}")
except Exception as e:
    print(f"Error al cargar el archivo CSV: {e}")
    raise

datos.fillna('No aplica', inplace=True)

# Seleccionar solo la primera fila y crear el diccionario de reemplazos
primera_fila = datos.iloc[0]
replacements = {f"#{col}#": str(primera_fila[col]) for col in datos.columns}
print(f"Reemplazos para la primera fila: {replacements}")

def replace_placeholder_in_paragraph(paragraph, replacements):
    for key, value in replacements.items():
        if key in paragraph.text:
            # Encontrar todos los runs que contienen el marcador
            # y combinar los runs si es necesario
            inline = paragraph.runs
            # Crear una lista de índices de runs a combinar
            indices = []
            for i in range(len(inline)):
                if key in inline[i].text:
                    indices.append(i)
            if indices:
                # Combinar los textos de los runs que contienen el marcador
                combined_text = ''.join([inline[i].text for i in indices])
                # Reemplazar el marcador en el texto combinado
                replaced_text = combined_text.replace(key, value)
                # Borrar el texto de los runs originales
                for i in indices:
                    inline[i].text = ''
                # Asignar el texto reemplazado al primer run
                inline[indices[0]].text = replaced_text
                # Copiar el formato del primer run a los demás si es necesario
                for i in indices[1:]:
                    inline[i].text = ''

def replace_text_in_doc(doc, replacements):
    # Reemplazar en el cuerpo del documento
    for paragraph in doc.paragraphs:
        replace_placeholder_in_paragraph(paragraph, replacements)
    # Reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholder_in_paragraph(paragraph, replacements)
    # Reemplazar en cabeceras y pies de página
    for section in doc.sections:
        # Cabecera
        header = section.header
        for paragraph in header.paragraphs:
            replace_placeholder_in_paragraph(paragraph, replacements)
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_placeholder_in_paragraph(paragraph, replacements)
        # Pie de página
        footer = section.footer
        for paragraph in footer.paragraphs:
            replace_placeholder_in_paragraph(paragraph, replacements)
        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_placeholder_in_paragraph(paragraph, replacements)

# Crear el documento de salida
output_path = os.path.join(output_folder, "muestra.docx")
try:
    doc = Document(plantilla)
    replace_text_in_doc(doc, replacements)
    doc.save(output_path)
    print(f"Carta generada correctamente: {output_path}")
except Exception as e:
    print(f"Error al procesar la plantilla .docx: {e}")
    raise
